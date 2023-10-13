# coding=utf-8
import os
import cx_Oracle 
import pandas as pd
import re
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Pt, Length
from docx.oxml import parse_xml
from docx.enum.table import WD_ROW_HEIGHT_RULE
os.environ['NLS_LANG'] = '.UTF8'


# 数据库连接信息
username = ''
password = ''
host = ''
dsn = cx_Oracle.makedsn('', '', service_name='orcl') 
conn = cx_Oracle.connect(user='', password='', dsn=dsn)
# 创建游标
cur = conn.cursor()


# 执行查询
query = "SELECT XNXQDM,KCM,JSM,XKZRS,ZC,SKXQ,KSJC,JSJC,JASMC FROM username.JSKB"
query1 = "SELECT XNXQDM,ZC,XQ,KSJC,JSJC,JASMC FROM username.CLASSROOM_LEND"
df = pd.read_sql(query, conn)
df = df[df['XNXQDM'] == '2023-2024-1'] 
df = df[df['JASMC'].str.startswith('计算中心', na=False)]

df1 = pd.read_sql(query1, conn)
df1 = df1[df1['XNXQDM'] == '2023-2024-1'] 
df1 = df1[df1['JASMC'].str.startswith('计算中心', na=False)]



def get_week(w):
    
    # 1 w 5-6,8-11周
    # 2 w 5-6,8周
    # 3 w 11周
    # 4 w 5,11周
    # 5 w 5-6周
    # 6 w 8
    # 7 w 8,9,10
    # 8 w ''
    # weeks [5,6,8,9,10,11]
    weeks = []
    if ' ' in w:                                           # 8
        weeks.append(' ')
    elif "周" not in w and "," not in w and "-" not in w:  # 6
            return [int(w)]
    # elif "周" not in w and "," in w and "-" not in w:    # 7
    #     start, end = map(int, w.split(','))
    #     weeks.extend(range(start, end + 1))
    elif "-" not in w and "," not in w:                    # 3
        return [    int(    w[:-1]   )   ]
    elif "," not in w:                                     # 5
        start, end = map(int, w[:-1].split('-'))
        weeks.extend(range(start, end + 1))
    else:                                                  # 1,2,4
        sub_list= w[:-1].split(',')                     
        for sub in sub_list:    
            if "-" in sub:                                 # 1,2
                start, end = map(int, sub.split('-'))
                weeks.extend(range(start, end + 1))
            else:                                          # 4 
                weeks.append(int(sub))

    return weeks



# 创建新的Word文档
doc = Document()
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.styles['Normal'].font.size = Pt(10.5)
doc.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
from docx.oxml.ns import qn
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
want_week =int(input("请输入要查询第几周数据:"))

#调节标题字体,颜色,大小
Head = doc.add_heading("",level=1)# 这里不填标题内容 
Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run  = Head.add_run("哈尔滨工程大学 教室课表")
run.font.name = u'宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run.font.size = Pt(14)
run.bold = True
run.font.color.rgb = RGBColor(0,0,0)
Head1 = doc.add_heading("",level=1)# 这里不填标题内容 
Head1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run1  = Head1.add_run("学年学期:2023-2024学年1学期  第{}周  所属功能区：计算机房  所属教学楼：机房  座位数:64".format(want_week))
run1.font.name = u'宋体'
run1._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run1.font.size = Pt(9)
run1.font.color.rgb = RGBColor(0,0,0)

# 遍历查询结果并将数据写入表格
table1 = doc.add_table(rows=6, cols=8)
table1.style = 'Table Grid' #设置表格样式

# 设置列宽
for i, row in enumerate(table1.rows):
    for j, cell in enumerate(row.cells):
        if j == 0:
            cell.width = Cm(2)
        else:
            cell.width = Cm(3.19)
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:vAlign xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="center"/>'))  # 垂直居中
# 写入表头  
header_cells1 = table1.rows[0].cells  
table1.cell(0,0).text = '星期节次'  
table1.cell(0,1).text = '星期一'  
table1.cell(0,2).text = '星期二'  
table1.cell(0,3).text = '星期三'  
table1.cell(0,4).text = '星期四'  
table1.cell(0,5).text = '星期五'  
table1.cell(0,6).text = '星期六'  
table1.cell(0,7).text = '星期日'  
table1.cell(1,0).text = '第一大节'  
table1.cell(2,0).text = '第二大节'  
table1.cell(3,0).text = '第三大节'  
table1.cell(4,0).text = '第四大节'  
table1.cell(5,0).text = '第五大节'
# 设置表头第一列宽度  
table1.columns[0].width = Cm(2)  
  
# 设置第一行高度  
table1.rows[0].height = Cm(0.64)  
table1.rows[0].height_rule=WD_ROW_HEIGHT_RULE.EXACTLY
table1.rows[1].height = Cm(2.52)  
table1.rows[1].height_rule=WD_ROW_HEIGHT_RULE.EXACTLY
table1.rows[2].height = Cm(2.52)  
table1.rows[2].height_rule=WD_ROW_HEIGHT_RULE.EXACTLY
table1.rows[3].height = Cm(2.52)  
table1.rows[3].height_rule=WD_ROW_HEIGHT_RULE.EXACTLY
table1.rows[4].height = Cm(2.52)  
table1.rows[4].height_rule=WD_ROW_HEIGHT_RULE.EXACTLY
table1.rows[5].height = Cm(2.52)  
table1.rows[5].height_rule=WD_ROW_HEIGHT_RULE.EXACTLY
        
#创建表头字典用于比对
weekday = {'星期一':1,'星期二':2,'星期三':3,'星期四':4,'星期五':5,'星期六':6,'星期日':7}
lesson = {
    '1':1,
    "2":1,
    "3":2,
    "4":2,
    "5":2,
    "6":3,
    "7":3,
    "8":4,
    "9":4,
    "10":4,
    "11":5,
    "12":5,
    "13":5,
}

# result_lists = []  # 存储多个列表的列表
# for i, row in df.iterrows():
#     list_week = list(get_all_weeks(row['ZC']))                        #list_week里存放的是需要的数据的周，以列表的形式
#     result_lists.append(list_week)                                    # 将当前循环中生成的列表添加到结果列表中
# #创建一个空字符串
# result = ''

# # 遍历df的每一行
# for index, row in df.iterrows():
#     # 使用apply方法逐元素拼接
#     row_str = row.drop('XNXQDM').apply(str).str.cat(sep=',')
#     # 将拼接后的行添加到结果字符串中
#     result += row_str + '\n'
# # print(result)


for index, row in df.iterrows():
    # print(row['SKXQ'],row['KSJC'])
    # 5-6,8-11周
    weeks = get_week(row['ZC'])
    if want_week in weeks: 
        # print("这节课在期望周存在",row["KCM"],weeks)
        col_index=weekday[row['SKXQ']]
        row_index=lesson[row['KSJC']]
        cell = table1.cell(row_index,col_index)               
        cell.text += ('\n' + ','.join(filter(None, [row['KCM'],
                                                        row['JSM'],
                                                        f"{str(row['XKZRS'])}人",
                                                        str(row['ZC']),
                                                        str(row['SKXQ']),
                                                        f"{row['KSJC']}-{row['JSJC']}节",
                                                        row['JASMC']
                                                        ])))
        
        if (int(row['JSJC'])-2) > (int(row['KSJC'])+1):
            row_index=lesson[row['JSJC']]
            cell = table1.cell(row_index,col_index)               
            cell.text += ('\n' + ','.join(filter(None, [row['KCM'],
                                                            row['JSM'],
                                                            f"{str(row['XKZRS'])}人",
                                                            str(row['ZC']),
                                                            str(row['SKXQ']),
                                                            f"{row['KSJC']}-{row['JSJC']}节",
                                                            row['JASMC']
                                                            ])))

for index,row in df1.iterrows():
    print(row['XQ'],row['KSJC'],type(row['XQ']),type(row['KSJC']),int(row['XQ']),type(int(row['XQ'])))
    weeks = get_week(row['ZC'])
    if want_week in weeks:
    
        col_index=int(row['XQ']) 
        row_index=lesson[row['KSJC']]
        cell = table1.cell(row_index,col_index) 
        cell.text += ('\n' + ','.join(filter(None, [f"{str(row['ZC'])}周",
                                                            f"星期{ str(row['XQ'])}",
                                                            f"{row['KSJC']}-{row['JSJC']}节",
                                                            row['JASMC']
                                                            ])))+ "(" + "借用" + ")"
        if (int(row['JSJC'])-2) > (int(row['KSJC'])+1):
            row_index=lesson[row['JSJC']]
            cell = table1.cell(row_index,col_index) 
            cell.text += ('\n' + ','.join(filter(None, [f"{str(row['ZC'])}周",
                                                                f"星期{ str(row['XQ'])}",
                                                                f"{row['KSJC']}-{row['JSJC']}节",
                                                                row['JASMC']
                                                                ])))+ "(" + "借用" + ")"


# 其他设置和保存文档
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21)
section.left_margin = Cm(1.91)
section.right_margin = Cm(1.91)
section.top_margin = Cm(0.68)
section.bottom_margin = Cm(0.68)
#保存Word文档
doc.save('timetable2.docx')






